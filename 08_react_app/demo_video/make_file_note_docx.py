"""Generate the realistic SAR SharePoint seed doc (.docx) from the file note.
No pipeline-walkthrough meta box; includes one named third party (Mrs Sarah Quinn)
so AI_CLASSIFY flags third-party data. Default black-on-white styling — a genuine
council record. Run: python3 make_file_note_docx.py"""
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

OUT = "audit/demo-assets/2026-04-02_ASC-2026-04021_file_note.docx"
BLUE = RGBColor(0x1F, 0x4E, 0x79)

doc = Document()
normal = doc.styles["Normal"]
normal.font.name = "Calibri"
normal.font.size = Pt(11)

def heading(text, size, color=None, space_before=10, space_after=4):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after = Pt(space_after)
    r = p.add_run(text)
    r.bold = True
    r.font.size = Pt(size)
    if color:
        r.font.color.rgb = color
    return p

def labelled(label, value):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(2)
    r = p.add_run(label)
    r.bold = True
    p.add_run(value)
    return p

def body(text):
    p = doc.add_paragraph(text)
    p.paragraph_format.space_after = Pt(8)
    return p

# Title
heading("Exampleton Borough Council — Adult Social Care & Housing Options", 18, BLUE, space_before=0, space_after=8)

# Metadata block
labelled("Classification: ", "OFFICIAL — SENSITIVE")
labelled("Record type: ", "Temporary accommodation / social care file note")
labelled("Data subject: ", "Mr James Whitfield")
labelled("Reference: ", "ASC-2026-04021    Note created: 2 April 2026")

# File note
heading("File note", 13, BLUE)
body(
    "Mr James Whitfield contacted the Housing Options team on 2 April 2026 regarding his "
    "temporary accommodation placement following the ongoing Housing Benefit overpayment "
    "dispute (claim HB-2026-55821). Mr Whitfield reports that the temporary accommodation "
    "is unsuitable for his mobility needs and has asked that Adult Social Care review his "
    "support plan."
)
body(
    "An initial social care needs assessment has been scheduled. Mr Whitfield's neighbour, "
    "Mrs Sarah Quinn (contactable on 0117 900 4412), provided a supporting statement "
    "describing the accessibility problems she has witnessed at the property and offered to "
    "speak with the assessing officer."
)
body(
    "Mr Whitfield also wishes this note, and the earlier complaint he raised, to be included "
    "in any Subject Access Request response so that he has a complete picture of the records "
    "the council holds about him. Third-party personal data recorded here, including Mrs "
    "Quinn's name and contact number, must be removed before any SAR disclosure to Mr "
    "Whitfield (s.40 DPA 2018)."
)

# Footer disclaimer
doc.add_paragraph().paragraph_format.space_after = Pt(4)
p = doc.add_paragraph()
r = p.add_run(
    "SYNTHETIC TRAINING DOCUMENT — contains fabricated personal data for demonstration only. "
    "Not a real person or case."
)
r.italic = True
r.font.size = Pt(9)
r.font.color.rgb = RGBColor(0x70, 0x70, 0x70)

doc.save(OUT)
print("WROTE", OUT)
